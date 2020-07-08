#!/usr/bin/env python3.8

import sys
# import string
from lib import logger
from docx import Document
# from docx.shared import Inches
from colorama import Fore, Back, Style

class Reporting :

    def __init__( self, db_name, rep_name, log ) :
        self.stash = logger.Stash( db_name, log )
        self.log = log
        self.db_name = db_name

        self.rep_file = rep_name
        self.doc = Document()

    def na( self, value ):
        return value if value else 'NA'

    def progress(self, count, total, status=''):
        bar_len = 60
        filled_len = int(round(bar_len * count / float(total)))

        percents = round(100.0 * count / float(total), 1)
        bar = Fore.GREEN + '=' * filled_len + Fore.RED + '-' * (bar_len - filled_len)

        sys.stdout.write( f'{Fore.BLUE}{Style.DIM}[{bar}{Fore.BLUE}] {percents}% --> {status}{Fore.RESET}{Style.RESET_ALL}{" "*10}\r')
        sys.stdout.flush()


    def writeReport( self ) :
        self.log.info( f'Retrieving data from {self.db_name}' )

        target = self.stash.get_column( 'subdomains', 'domain' )

        subdom = self.stash.get_column( 'whois', 'group_concat( hostname , char(10) ), ip', grB='ip' )
        # subdom = [(subd , ip),(),()]

        ip_asn = self.stash.get_column( 'server_info', 'asn, group_concat( ip, char(10)), organization, coordinate, isp', grB='asn' )
        # ip_asn = [(asn , ip, organization, coordinate, isp),(),()]

        asn_info = self.stash.get_column( 'asn_info', '*', grB='asn' )
        # yes I know, if you don't, just move on. Nothing to see here.
        # asn_info = [(asn , country, registry, cidr, description, registration date),(),()]

        services = self.stash.get_column( 'services', '''ip, group_concat( 'Port: ' || port || char(10) || 'Service: ' || service )''', grB='ip' )
        # services = [(ip , service/port),(),()]

        employees = self.stash.get_column( 'employees', 'name, position, prof_link, email' , grB='name' )
        linkedin_link = self.stash.get_column( 'employees', 'comp_link' )
        # employees = [(name, position, prof_link, prof_pic, email),(),()]

        potential_employees = self.stash.get_column( 'potential_employees', 'name, position, prof_link, email' , grB='name' )
        linkedin_link = self.stash.get_column( 'potential_employees', 'comp_link' )
        # potential_employees = [(name, position, prof_link, prof_pic, email),(),()]



        self.log.info( f'Generating the report {self.rep_file}' )
        self.doc.add_heading( 'OSINT report' , 0 )
        self.doc.add_paragraph( 'This report has been automatically generated with teh LazyOSINT framework.\nIt will present all the data retreived divided into tables.' )

        self.doc.add_heading('Table of Subdomains and relative IPs', level=1)
        subdT = self.doc.add_table( rows=1, cols=2 )
        hdr_cells = subdT.rows[0].cells
        hdr_cells[0].text = 'IP'
        hdr_cells[1].text = 'Subdomain'
        s = 0
        l = len(subdom)
        for sub, ip in subdom:
            row_cells = subdT.add_row().cells
            row_cells[0].text = self.na(ip)
            row_cells[1].text = self.na(sub).replace( '<BR>', '\n' )
            self.progress( s, l, 'Writing all Subdomains info.' )
            s += 1


        self.doc.add_heading('Table of information about the Servers', level=1)
        subdIS = self.doc.add_table( rows=1, cols=5 )
        hdr_cells = subdIS.rows[0].cells
        hdr_cells[0].text = 'ASN'
        hdr_cells[1].text = 'IPs'
        hdr_cells[2].text = 'Organization'
        hdr_cells[3].text = 'Coordinate'
        hdr_cells[4].text = 'ISP'
        s = 0
        l = len(ip_asn)
        for asn, ip, org, coo, isp in ip_asn:
            row_cells = subdIS.add_row().cells
            row_cells[0].text = self.na(asn)
            row_cells[1].text = self.na(ip)
            row_cells[2].text = self.na(org)
            row_cells[3].text = self.na(coo)
            row_cells[4].text = self.na(isp)
            self.progress( s, l, 'Writing all server info.' )
            s += 1


        # asn_info = [(asn , country, registry, cidr, description, registration date),(),()]
        self.doc.add_heading('Table of ASN and relative info', level=1)
        subdIA = self.doc.add_table( rows=1, cols=6 )
        hdr_cells = subdIA.rows[0].cells
        hdr_cells[0].text = 'ASN'
        hdr_cells[1].text = 'Country'
        hdr_cells[2].text = 'Registry'
        hdr_cells[3].text = 'CIDR'
        hdr_cells[4].text = 'Description'
        hdr_cells[5].text = 'Reg. date'
        s = 0
        l = len(asn_info)
        for asn, cou, reg, cidr, desc, dat in asn_info:
            row_cells = subdIA.add_row().cells
            row_cells[0].text = self.na(asn)
            row_cells[1].text = self.na(cou)
            row_cells[2].text = self.na(reg)
            row_cells[3].text = self.na(cidr)
            row_cells[4].text = self.na(desc)
            row_cells[5].text = self.na(dat)
            self.progress( s, l, 'Writing all ASN related info.' )
            s += 1


        # services = [(ip , service/port),(),()]
        self.doc.add_heading('Table of Services for each IP', level=1)
        subdP = self.doc.add_table( rows=1, cols=2 )
        hdr_cells = subdP.rows[0].cells
        hdr_cells[0].text = 'IP'
        hdr_cells[1].text = 'Ports'
        s = 0
        l = len(services)
        for ip, p in services:
            row_cells = subdP.add_row().cells
            row_cells[0].text = self.na(ip)
            row_cells[1].text = self.na(p).replace( '\r' , '' ).replace( ',Port' , 'Port' )
            self.progress( s, l, 'Writing all info about open ports.' )
            s += 1


        # employees = [(name, position, prof_link, prof_pic, email),(),()]
        self.doc.add_heading('Table of information about Employees', level=1)
        self.doc.add_paragraph( 'The contents of the following table are retreived from LinkedIn.\nThe emails are generated using a sample.' )
        subdE = self.doc.add_table( rows=1, cols=4 )
        hdr_cells = subdE.rows[0].cells
        hdr_cells[0].text = 'Full Name'
        hdr_cells[1].text = 'Position'
        hdr_cells[2].text = 'Profile Link'
        hdr_cells[3].text = 'Probable Email'
        s = 0
        l = len( employees )
        for name, pos, prof, email in employees:
            row_cells = subdE.add_row().cells
            row_cells[0].text = self.na(name)
            row_cells[1].text = self.na(pos)
            row_cells[2].text = self.na(prof)
            row_cells[3].text = self.na(email)
            self.progress( s, l, 'Writing all employees data.' )
            s += 1

        # employees = [(name, position, prof_link, prof_pic, email),(),()]
        self.doc.add_heading('Table of information about Potential Employees', level=1)
        self.doc.add_paragraph( 'The contents of the following table are retreived from Google dorking.\nThe emails are generated using a sample.' )
        subdE = self.doc.add_table( rows=1, cols=4 )
        hdr_cells = subdE.rows[0].cells
        hdr_cells[0].text = 'Full Name'
        hdr_cells[1].text = 'Position'
        hdr_cells[2].text = 'Profile Link'
        hdr_cells[3].text = 'Probable Email'
        s = 0
        l = len( potential_employees )
        for name, pos, prof, email in potential_employees:
            row_cells = subdE.add_row().cells
            row_cells[0].text = self.na(name)
            row_cells[1].text = self.na(pos)
            row_cells[2].text = self.na(prof)
            row_cells[3].text = self.na(email)
            self.progress( s, l, 'Writing all potential employees data.' )
            s += 1

        self.doc.save( self.rep_file )

